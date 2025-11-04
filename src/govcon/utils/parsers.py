"""Document parsing utilities for knowledge base."""

from __future__ import annotations

from pathlib import Path
from typing import Optional

from govcon.utils.logger import get_logger

logger = get_logger(__name__)


def parse_document(file_path: str) -> str:
    """
    Parse a document and extract text content.

    Args:
        file_path: Path to the document file

    Returns:
        Extracted text content

    Raises:
        ValueError: If file type is not supported
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    file_type = path.suffix.lower()

    if file_type == ".txt":
        return parse_text(file_path)
    elif file_type == ".md":
        return parse_markdown(file_path)
    elif file_type == ".pdf":
        return parse_pdf(file_path)
    elif file_type in [".docx", ".doc"]:
        return parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def parse_text(file_path: str) -> str:
    """Parse plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def parse_markdown(file_path: str) -> str:
    """Parse markdown file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def parse_pdf(file_path: str) -> str:
    """
    Parse PDF file and extract text.

    Args:
        file_path: Path to PDF file

    Returns:
        Extracted text content
    """
    try:
        import PyPDF2

        text_content = []
        with open(file_path, "rb") as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {e}")
                    continue

        return "\n\n".join(text_content)

    except ImportError:
        logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
        raise
    except Exception as e:
        logger.error(f"Failed to parse PDF {file_path}: {e}")
        raise


def parse_docx(file_path: str) -> str:
    """
    Parse DOCX file and extract text.

    Args:
        file_path: Path to DOCX file

    Returns:
        Extracted text content
    """
    try:
        import docx

        doc = docx.Document(file_path)
        text_content = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_content.append(row_text)

        return "\n\n".join(text_content)

    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        raise
    except Exception as e:
        logger.error(f"Failed to parse DOCX {file_path}: {e}")
        raise


def chunk_text(
    text: str,
    chunk_size: int = 1000,
    overlap: int = 200,
    min_chunk_size: Optional[int] = None,
) -> list[str]:
    """
    Split text into overlapping chunks for vector storage.

    Args:
        text: Text to chunk
        chunk_size: Target size of each chunk in characters
        overlap: Number of characters to overlap between chunks
        min_chunk_size: Minimum chunk size (default: chunk_size // 2)

    Returns:
        List of text chunks
    """
    if min_chunk_size is None:
        min_chunk_size = chunk_size // 2

    # Split by paragraphs first
    paragraphs = text.split("\n\n")

    chunks = []
    current_chunk = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # If adding this paragraph exceeds chunk size
        if len(current_chunk) + len(para) > chunk_size:
            # Save current chunk if it's not empty
            if len(current_chunk) >= min_chunk_size:
                chunks.append(current_chunk.strip())
                # Start new chunk with overlap from previous
                current_chunk = current_chunk[-overlap:] + "\n\n" + para
            else:
                # Current chunk is too small, add paragraph anyway
                current_chunk += "\n\n" + para
        else:
            # Add paragraph to current chunk
            if current_chunk:
                current_chunk += "\n\n" + para
            else:
                current_chunk = para

    # Add final chunk
    if len(current_chunk) >= min_chunk_size:
        chunks.append(current_chunk.strip())
    elif chunks:
        # Merge with last chunk if too small
        chunks[-1] += "\n\n" + current_chunk
    else:
        # No chunks yet, add it anyway
        chunks.append(current_chunk.strip())

    return [c for c in chunks if c.strip()]


def extract_metadata(text: str, file_name: str) -> dict:
    """
    Extract metadata from document text.

    Args:
        text: Document text
        file_name: Original file name

    Returns:
        Dictionary with extracted metadata
    """
    metadata = {
        "file_name": file_name,
        "char_count": len(text),
        "word_count": len(text.split()),
        "line_count": len(text.splitlines()),
    }

    # Try to extract title (first non-empty line)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if lines:
        # Check if first line looks like a title (short, no period at end)
        first_line = lines[0]
        if len(first_line) < 200 and not first_line.endswith("."):
            metadata["extracted_title"] = first_line

    # Check for common section headers
    sections = []
    common_headers = [
        "EXECUTIVE SUMMARY",
        "TECHNICAL APPROACH",
        "MANAGEMENT APPROACH",
        "PAST PERFORMANCE",
        "PRICING",
        "SCOPE OF WORK",
        "QUALIFICATIONS",
    ]

    for header in common_headers:
        if header in text.upper():
            sections.append(header.lower())

    if sections:
        metadata["sections_found"] = sections

    return metadata
